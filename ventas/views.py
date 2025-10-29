# ventas/views.py
"""
Maneja la lógica de las vistas para la aplicación 'ventas'.

Incluye:
- CRUD para Órdenes de Compra (Crear, Listar, Ver Detalle).
- Validación de stock y actualización atómica durante la creación.
- Generación de documentos (PDF y DOCX) para las órdenes.
"""

# --- Importaciones de Django ---
from django.shortcuts import render, get_object_or_404, redirect, reverse
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction # Para asegurar la integridad de la BD
from django.template.loader import render_to_string
from django.contrib.humanize.templatetags.humanize import intcomma
from django.contrib.staticfiles import finders # Para encontrar el logo
import os

# --- Importaciones de Modelos y Forms ---
from .models import OrdenCompra, DetalleOrden
from inventario.models import Producto
from .forms import OrdenCompraForm, DetalleOrdenFormSet

# --- Importaciones de Librerías (PDF, DOCX, etc.) ---
from io import BytesIO # Buffer en memoria para archivos

# --- IMPORTACIONES PARA PDF (ReportLab) ---
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.units import mm # Para usar milímetros
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, Image, Frame, BaseDocTemplate, PageTemplate
from reportlab.platypus.flowables import HRFlowable # Línea horizontal
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# --- IMPORTACIONES PARA DOCX (python-docx) ---
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Vistas de Órdenes de Compra ---

@login_required
def lista_ordenes(request):
    """
    Muestra una lista de todas las órdenes de compra registradas,
    ordenadas por fecha descendente.
    """
    # .prefetch_related() optimiza la consulta al traer los detalles
    # y productos relacionados en una sola consulta adicional.
    ordenes = OrdenCompra.objects.prefetch_related('detalles__producto').all().order_by('-fecha')
    return render(request, 'ventas/lista_ordenes.html', {'ordenes': ordenes})


@login_required
@transaction.atomic # Decorador clave: Si algo falla, revierte todos los cambios en la BD.
def crear_orden(request):
    """
    Maneja la creación de una nueva Orden de Compra (GET y POST).

    En POST:
    1. Valida el formulario de la orden y el formset de detalles.
    2. Itera los detalles para validar el stock de cada producto.
    3. Si el stock es válido, calcula el total.
    4. Guarda la orden y sus detalles.
    5. Disminuye el stock de los productos vendidos.
    Todo esto ocurre dentro de una transacción atómica.
    """
    if request.method == 'POST':
        orden_form = OrdenCompraForm(request.POST)
        # El prefix='detalles' debe coincidir con el usado en el template
        detalle_formset = DetalleOrdenFormSet(request.POST, prefix='detalles')

        if orden_form.is_valid() and detalle_formset.is_valid():
            try:
                # 1. Preparar datos (sin guardar en BD aún)
                orden = orden_form.save(commit=False)
                detalles_instancias = detalle_formset.save(commit=False)
                
                total_orden = 0
                productos_a_actualizar = [] # Lista para guardar {producto_db, cantidad}
                detalles_a_guardar = [] # Lista para guardar instancias de detalle
                valid_details_count = 0 # Contador de líneas válidas

                # 2. Validar stock y calcular total
                for form in detalle_formset:
                    # Ignorar formularios vacíos o marcados para borrar
                    if form.cleaned_data and not form.cleaned_data.get('DELETE', False):
                        producto = form.cleaned_data.get('producto')
                        cantidad = form.cleaned_data.get('cantidad')
                        precio = form.cleaned_data.get('precio_unitario')

                        # Asegurarse que la línea tiene todos los datos
                        if producto and cantidad and cantidad > 0 and precio is not None:
                            valid_details_count += 1
                            detalle = form.instance
                            
                            # ¡CRÍTICO! Bloquear la fila del producto para evitar
                            # que dos ventas del mismo producto ocurran a la vez (race condition).
                            producto_db = Producto.objects.select_for_update().get(pk=producto.pk)
                            
                            # Validar stock
                            if producto_db.stock < cantidad:
                                # Si falla, se lanza una excepción que será capturada
                                # y la transacción se revertirá.
                                raise Exception(f"No hay suficiente stock para: {producto.nombre} (disponible: {producto_db.stock})")

                            # Si el stock es válido, acumular
                            total_linea = cantidad * precio
                            total_orden += total_linea
                            detalles_a_guardar.append(detalle)
                            productos_a_actualizar.append({'producto': producto_db, 'cantidad': cantidad})

                # Validar que al menos se añadió un producto
                if valid_details_count == 0:
                    raise Exception("Debes añadir al menos un producto válido a la orden.")

                # 3. Guardar en Base de Datos (si todo fue válido)
                orden.total = total_orden
                orden.save() # Guardar la orden principal (obtiene un ID)

                # 4. Guardar detalles y actualizar stock
                for detalle in detalles_a_guardar:
                    detalle.orden = orden # Asignar la orden ya guardada
                    detalle.save()

                for item in productos_a_actualizar:
                    # Usar el método del modelo para disminuir stock
                    item['producto'].disminuir_stock(item['cantidad'])

                detalle_formset.save_m2m() # Guardar relaciones (aunque aquí no hay m2m)

                messages.success(request, f"Orden {orden.numero_venta} creada exitosamente.")
                return redirect('ventas:detalle_orden', orden_id=orden.pk)

            except Exception as e:
                # Si algo falló (ej. stock), se muestra el error
                # y @transaction.atomic revierte la creación de la orden.
                messages.error(request, f"Error al crear la orden: {e}")

    else: # Método GET
        # Mostrar formularios vacíos
        orden_form = OrdenCompraForm()
        detalle_formset = DetalleOrdenFormSet(prefix='detalles')

    context = {
        'orden_form': orden_form,
        'detalle_formset': detalle_formset
    }
    return render(request, 'ventas/crear_orden.html', context)


@login_required
def detalle_orden(request, orden_id):
    """
    Muestra la vista de detalle (ticket) de una orden de compra específica.
    """
    orden = get_object_or_404(OrdenCompra.objects.prefetch_related('detalles__producto'), pk=orden_id)
    return render(request, 'ventas/detalle_orden.html', {'orden': orden})


# --- Vistas de Descarga de Documentos ---

@login_required
def descargar_orden_pdf(request, orden_id):
    """
    Genera y sirve un archivo PDF de la orden de compra usando ReportLab.
    El formato está diseñado para simular un ticket de 80mm.
    """
    orden = get_object_or_404(OrdenCompra.objects.prefetch_related('detalles__producto'), pk=orden_id)

    # 1. Configurar el buffer en memoria
    buffer = BytesIO()

    # 2. Configuración del Documento (Ticket 80mm)
    ticket_width = 80 * mm
    ticket_height = 200 * mm # Altura estimada, puede crecer
    pagesize = (ticket_width, ticket_height)
    margin = 5 * mm
    effective_width = ticket_width - 2 * margin # Ancho útil

    doc = BaseDocTemplate(buffer, pagesize=pagesize,
                          leftMargin=margin, rightMargin=margin,
                          topMargin=margin, bottomMargin=margin)

    # 3. Definición de Estilos de Párrafo
    styles = getSampleStyleSheet()
    style_base = ParagraphStyle(name='Base', parent=styles['Normal'], fontSize=8, leading=10)
    style_normal = ParagraphStyle(name='Normal', parent=style_base, alignment=TA_LEFT)
    style_bold = ParagraphStyle(name='Bold', parent=style_base, fontName='Helvetica-Bold')
    style_header_info = ParagraphStyle(name='HeaderInfo', parent=style_base, fontSize=7, leading=8.5, alignment=TA_LEFT)
    style_header_name = ParagraphStyle(name='HeaderName', parent=style_header_info, fontName='Helvetica-Bold')
    style_order_title = ParagraphStyle(name='OrderTitle', parent=style_bold, fontSize=9, leading=11, alignment=TA_LEFT)
    style_center = ParagraphStyle(name='Center', parent=style_base, alignment=TA_CENTER)
    style_right_bold = ParagraphStyle(name='RightBold', parent=style_bold, alignment=TA_RIGHT)
    style_total_label = ParagraphStyle(name='TotalLabel', parent=style_right_bold, fontSize=10, leading = 12)
    style_gracias = ParagraphStyle(name='Gracias', parent=style_center, fontSize=8, leading=10)
    style_table_header = ParagraphStyle(name='TableHeader', parent=style_bold, fontSize=7, alignment=TA_CENTER)
    style_table_cell = ParagraphStyle(name='TableCell', parent=style_base, fontSize=7)
    style_table_cell_right = ParagraphStyle(name='TableCellRight', parent=style_table_cell, alignment=TA_RIGHT)
    style_table_product = ParagraphStyle(name='TableProduct', parent=style_table_cell, alignment=TA_LEFT)

    # 4. Contenido del PDF (El "Story" de ReportLab)
    story = []

    # --- Encabezado con Logo ---
    logo_path_relative = 'app/img/logo.png' # Ruta en /core/static/
    logo_path_absolute = finders.find(logo_path_relative)

    if logo_path_absolute and os.path.exists(logo_path_absolute):
        img = Image(logo_path_absolute, height=15*mm, width=15*mm) # Ajustar tamaño
        img.hAlign = 'LEFT'
        
        # Usar una tabla para alinear logo e info de la empresa
        header_data = [
            [img, Paragraph("<b>CONSTRUCCIONES V & G<br/>LIZ CASTILLO GARCIA SPA</b><br/>"
                           "RUT: 77.858.577-4<br/>"
                           "Dirección: Vilaco 301, Toconao<br/>"
                           "Teléfono: +56 9 52341652", style_header_info)]
        ]
        header_table = Table(header_data, colWidths=[20*mm, effective_width - 20*mm])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1*mm),
        ]))
        story.append(header_table)
    else:
        # Fallback si no se encuentra el logo
        story.append(Paragraph("CONSTRUCCIONES V & G LIZ CASTILLO GARCIA SPA", style_center))
        if not logo_path_absolute:
            messages.warning(request, f"No se encontró el archivo del logo en: {logo_path_relative}")

    story.append(Spacer(1, 1 * mm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black, spaceBefore=1*mm, spaceAfter=1*mm))

    # --- Info de la Orden ---
    story.append(Paragraph(f"Orden de Compra #{orden.numero_venta or orden.id}", style_order_title))
    story.append(Paragraph(f"Cliente: {orden.cliente}", style_normal))
    if orden.rut:
        story.append(Paragraph(f"RUT: {orden.rut}", style_normal))
    story.append(Paragraph(f"Fecha: {orden.fecha.strftime('%d-%m-%Y %H:%M')}", style_normal))
    if orden.direccion:
        story.append(Paragraph(f"Dirección: {orden.direccion}", style_normal))
    story.append(Spacer(1, 3 * mm))
    story.append(Paragraph("Detalle", style_bold))
    story.append(Spacer(1, 1 * mm))

    # --- Tabla de Productos ---
    headers = [
        Paragraph('Producto', style_table_header),
        Paragraph('Cant', style_table_header),
        Paragraph('P. Unitario', style_table_header),
        Paragraph('Total', style_table_header)
    ]
    data = [headers]
    detalles = orden.detalles.all()

    # Estilos de la tabla
    table_style_commands = [
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black), # Borde
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),    # Col 0 (Producto)
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),  # Col 1 (Cant)
        ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),  # Col 2 y 3 (Precios)
        ('LEFTPADDING', (0, 0), (-1, -1), 1.5*mm),
        ('RIGHTPADDING', (0, 0), (-1, -1), 1.5*mm),
    ]

    for detalle in detalles:
        data.append([
            Paragraph(detalle.producto.nombre, style_table_product),
            Paragraph(str(detalle.cantidad), style_table_cell),
            Paragraph(f"${intcomma(int(detalle.precio_unitario))}", style_table_cell_right),
            Paragraph(f"${intcomma(int(detalle.total_linea))}", style_table_cell_right)
        ])

    col_widths = [effective_width * 0.40, effective_width * 0.15, effective_width * 0.22, effective_width * 0.23]
    tabla = Table(data, colWidths=col_widths)
    tabla.setStyle(TableStyle(table_style_commands))
    story.append(tabla)
    story.append(Spacer(1, 3 * mm))

    # --- Total ---
    total_str = f"${intcomma(int(orden.total))}"
    total_data = [[Paragraph('Total a Pagar:', style_total_label), Paragraph(total_str, style_total_label)]]
    total_table = Table(total_data, colWidths=[effective_width * 0.6, effective_width * 0.4])
    total_table.setStyle(TableStyle([('ALIGN', (0, 0), (-1, -1), 'RIGHT')])) # Alinear todo a la derecha
    story.append(total_table)
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black, spaceBefore=1*mm, spaceAfter=1*mm))

    # --- Mensaje final ---
    story.append(Paragraph("¡Gracias por su compra!", style_gracias))
    story.append(Paragraph("Esperamos atenderle pronto.", style_gracias))

    # 5. Construir el PDF
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='normal')
    template = PageTemplate(id='main', frames=[frame])
    doc.addPageTemplates([template])

    try:
        doc.build(story)
    except Exception as e:
        print(f"Error al construir PDF con ReportLab: {e}")
        messages.error(request, f"Error al generar PDF: {e}")
        return redirect('ventas:detalle_orden', orden_id=orden.pk)

    # 6. Crear y devolver la Respuesta HTTP
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="orden_{orden.numero_venta or orden.id}.pdf"'
    return response


@login_required
def descargar_orden_docx(request, orden_id):
    """
    Genera y sirve un archivo DOCX (Word) de la orden de compra
    usando la librería python-docx.
    """
    orden = get_object_or_404(OrdenCompra.objects.prefetch_related('detalles__producto'), pk=orden_id)
    try:
        # 1. Crear documento y configurar márgenes
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # 2. Añadir contenido
        # --- Info Empresa ---
        p_empresa = document.add_paragraph()
        p_empresa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        runner = p_empresa.add_run(
            "CONSTRUCCIONES V & G LIZ CASTILLO GARCIA SPA\n"
            "RUT: 77.858.577-4\n"
            "Dirección: Vilaco 301, Toconao\n"
            "Teléfono: +56 9 52341652"
        )
        runner.font.size = Pt(8)
        runner.bold = True
        p_empresa.paragraph_format.space_after = Pt(0)

        document.add_paragraph("---" * 12).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Info Orden ---
        p_orden_info = document.add_paragraph()
        p_orden_info.add_run(f"Orden de Compra #{orden.numero_venta or orden.id}\n").bold = True
        p_orden_info.add_run(f"Cliente: {orden.cliente}\n")
        p_orden_info.add_run(f"Rut: {orden.rut or 'N/A'}\n")
        p_orden_info.add_run(f"Fecha: {orden.fecha.strftime('%d-%m-%Y %H:%M')}\n")
        p_orden_info.add_run(f"Dirección: {orden.direccion or 'N/A'}")
        for run in p_orden_info.runs:
            run.font.size = Pt(9)
        p_orden_info.paragraph_format.space_after = Pt(6)

        document.add_paragraph("---" * 12).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Tabla de Productos ---
        document.add_paragraph().add_run("Detalle").bold = True
        
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid' # Estilo con bordes
        table.autofit = False

        # Encabezados
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Producto'
        hdr_cells[1].text = 'Cant'
        hdr_cells[2].text = 'P. Unit.'
        hdr_cells[3].text = 'Total'

        # Datos
        for detalle in orden.detalles.all():
            row_cells = table.add_row().cells
            # Formatear números con separador de miles (usando intcomma)
            precio_unit_str = f"${intcomma(int(detalle.precio_unitario))}"
            total_linea_str = f"${intcomma(int(detalle.total_linea))}"

            row_cells[0].text = detalle.producto.nombre
            row_cells[1].text = str(detalle.cantidad)
            row_cells[2].text = precio_unit_str
            row_cells[3].text = total_linea_str
            
            # Alinear celdas numéricas
            for i, cell in enumerate(row_cells):
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                if i > 0:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Definir ancho de columnas
        table.columns[0].width = Inches(2.8)
        table.columns[1].width = Inches(0.5)
        table.columns[2].width = Inches(0.9)
        table.columns[3].width = Inches(1.0)

        # --- Total ---
        p_total = document.add_paragraph()
        p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_str = f"${intcomma(int(orden.total))}"
        runner_total = p_total.add_run(f"Total a Pagar: {total_str}")
        runner_total.bold = True
        runner_total.font.size = Pt(11)

        # 3. Guardar en buffer de memoria
        f = BytesIO()
        document.save(f)
        f.seek(0)

        # 4. Crear y devolver la Respuesta HTTP
        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="orden_{orden.numero_venta or orden.id}.docx"'
        f.close()
        return response

    except Exception as e:
        print(f"Error al generar DOCX: {e}")
        messages.error(request, f"Error al generar DOCX: {e}.")
        return redirect('ventas:detalle_orden', orden_id=orden.pk)