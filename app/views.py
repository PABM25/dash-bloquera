# app/views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse, reverse_lazy # <-- Añadir reverse_lazy
from django.http import HttpResponse, JsonResponse
from .models import *
from .forms import *
from io import BytesIO
from datetime import date
from django.db.models import Sum, F, Count
from django.db.models.functions import ExtractMonth, ExtractYear
import json
from django.db import transaction
from django.contrib import messages
from django.template.loader import render_to_string
import imgkit
import pdfkit 
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login
# --- Importaciones Modificadas/Añadidas ---
from django.contrib.auth.views import LoginView, PasswordChangeView, PasswordChangeDoneView 
from .forms import RegistroForm, EditProfileForm 
# --- FIN ---

# --- Vistas de Autenticación ---

class CustomLoginView(LoginView):
    template_name = 'app/login.html'
    redirect_authenticated_user = True

def register(request):
    if request.method == 'POST':
        form = RegistroForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            messages.success(request, '¡Registro exitoso! Has iniciado sesión.')
            return redirect('home')
    else:
        form = RegistroForm()
    return render(request, 'app/register.html', {'form': form})

# --- Vistas Home, Reporte, Lista Ordenes, Crear Orden ---

@login_required
def home(request):
    hoy = date.today()
    primer_dia_mes = hoy.replace(day=1)
    asistencia_del_mes = Asistencia.objects.filter(fecha__gte=primer_dia_mes).count()
    meses_etiquetas, ventas_mensuales, gastos_mensuales = reporte_graficos_data()
    datos_ventas_lista = [float(v.get('total_ventas', 0)) for v in ventas_mensuales]
    datos_gastos_lista = [float(g.get('total_gastos', 0)) for g in gastos_mensuales]
    total_ventas = sum(datos_ventas_lista)
    total_gastos = sum(datos_gastos_lista)
    total_general = total_ventas + total_gastos
    porcentaje_ventas = (total_ventas / total_general * 100) if total_general > 0 else 0
    porcentaje_gastos = (total_gastos / total_general * 100) if total_general > 0 else 0
    context = {
        'asistencia_del_mes': asistencia_del_mes,
        'total_ventas': total_ventas,
        'total_gastos': total_gastos,
        'meses_etiquetas_json': json.dumps(meses_etiquetas),
        'datos_ventas_json': json.dumps(datos_ventas_lista),
        'datos_gastos_json': json.dumps(datos_gastos_lista),
        'porcentaje_ventas': round(porcentaje_ventas, 1),
        'porcentaje_gastos': round(porcentaje_gastos, 1),
    }
    return render(request, 'app/home.html', context)

def reporte_graficos_data():
    ventas_qs = OrdenCompra.objects.annotate(mes=ExtractMonth('fecha'), ano=ExtractYear('fecha')).values('mes', 'ano').annotate(total_ventas=Sum('total')).order_by('ano', 'mes')
    gastos_qs = Gasto.objects.annotate(mes=ExtractMonth('fecha'), ano=ExtractYear('fecha')).values('mes', 'ano').annotate(total_gastos=Sum('monto')).order_by('ano', 'mes')
    meses_ventas = {f"{v['ano']}-{str(v['mes']).zfill(2)}" for v in ventas_qs}
    meses_gastos = {f"{g['ano']}-{str(g['mes']).zfill(2)}" for g in gastos_qs}
    meses_etiquetas = sorted(list(meses_ventas.union(meses_gastos)))
    ventas_dict = {f"{v['ano']}-{str(v['mes']).zfill(2)}": v for v in ventas_qs}
    gastos_dict = {f"{g['ano']}-{str(g['mes']).zfill(2)}": g for g in gastos_qs}
    ventas_final = [ventas_dict.get(mes, {'total_ventas': 0}) for mes in meses_etiquetas]
    gastos_final = [gastos_dict.get(mes, {'total_gastos': 0}) for mes in meses_etiquetas]
    return meses_etiquetas, ventas_final, gastos_final

@login_required
def lista_ordenes(request):
    ordenes = OrdenCompra.objects.prefetch_related('detalles__producto').all().order_by('-fecha')
    return render(request, 'app/lista_ordenes.html', {'ordenes': ordenes})

@login_required
def crear_orden(request):
    if request.method == 'POST':
        form = OrdenCompraForm(request.POST)
        formset = DetalleOrdenFormSet(request.POST)
        if form.is_valid() and formset.is_valid():
            try:
                with transaction.atomic():
                    orden = form.save(commit=False)
                    orden.save() 
                    detalles = formset.save(commit=False)
                    total_orden = 0
                    productos_a_guardar = []
                    valid_details_count = 0
                    for i, detalle_form in enumerate(formset):
                        if formset.can_delete and formset._should_delete_form(detalle_form): continue
                        if detalle_form.cleaned_data.get('producto') and detalle_form.cleaned_data.get('cantidad', 0) > 0:
                            valid_details_count += 1
                            detalle = detalle_form.instance
                            producto = detalle.producto
                            if not producto.disminuir_stock(detalle.cantidad): raise Exception(f"No hay suficiente stock para: {producto.nombre}")
                            total_linea = detalle.cantidad * detalle.precio_unitario
                            total_orden += total_linea
                            detalle.orden = orden 
                            productos_a_guardar.append(detalle)
                        elif detalle_form.has_changed() and not detalle_form.cleaned_data.get('producto'): pass
                    if valid_details_count == 0: raise Exception("Debes añadir al menos un producto válido a la orden.")
                    for detalle in productos_a_guardar: detalle.save()
                    orden.total = total_orden
                    orden.save(update_fields=['total']) 
                    messages.success(request, f"Orden {orden.numero_venta} creada exitosamente.")
                    return redirect('detalle_orden', orden_id=orden.pk)
            except Exception as e:
                messages.error(request, str(e))
                context = {'orden_form': form, 'detalle_formset': formset}
                return render(request, 'app/crear_orden.html', context)
    else:
        form = OrdenCompraForm()
        formset = DetalleOrdenFormSet()
    context = {'orden_form': form, 'detalle_formset': formset}
    return render(request, 'app/crear_orden.html', context)

@login_required
def detalle_orden(request, orden_id):
    orden = get_object_or_404(OrdenCompra.objects.prefetch_related('detalles__producto'), pk=orden_id)
    return render(request, 'app/detalle_orden.html', {'orden': orden})

# --- Vistas de Descarga ---
@login_required
def descargar_orden_pdf(request, orden_id):
    orden = get_object_or_404(OrdenCompra.objects.prefetch_related('detalles__producto'), pk=orden_id)
    html_string = render_to_string('app/detalle_orden_render.html', {'orden': orden, 'request': request})
    try:
        options = {'page-size': 'A7','margin-top': '5mm','margin-right': '5mm','margin-bottom': '5mm','margin-left': '5mm','encoding': "UTF-8",'enable-local-file-access': None,'quiet': ''}
        pdf_data = pdfkit.from_string(html_string, False, options=options)
        response = HttpResponse(pdf_data, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="orden_{orden.numero_venta or orden.id}.pdf"'
        return response
    except Exception as e:
        print(f"Error al generar PDF: {e}")
        return HttpResponse(f"Error al generar PDF: {e}", status=500)

@login_required
def descargar_orden_jpg(request, orden_id):
    orden = get_object_or_404(OrdenCompra.objects.prefetch_related('detalles__producto'), pk=orden_id)
    html_string = render_to_string('app/detalle_orden_render.html', {'orden': orden, 'request': request})
    try:
        options = {'format': 'jpg','encoding': "UTF-8",'quality': '90','enable-local-file-access': None,'quiet': ''}
        jpg_data = imgkit.from_string(html_string, False, options=options)
        response = HttpResponse(jpg_data, content_type='image/jpeg')
        response['Content-Disposition'] = f'attachment; filename="orden_{orden.numero_venta or orden.id}.jpg"'
        return response
    except Exception as e:
        print(f"Error al generar JPG: {e}")
        return HttpResponse(f"Error al generar JPG: {e}", status=500)

@login_required
def descargar_orden_docx(request, orden_id):
    orden = get_object_or_404(OrdenCompra.objects.prefetch_related('detalles__producto'), pk=orden_id)
    document = Document()
    sections = document.sections
    for section in sections: section.top_margin = Inches(0.4); section.bottom_margin = Inches(0.4); section.left_margin = Inches(0.5); section.right_margin = Inches(0.5)
    p_empresa = document.add_paragraph(); p_empresa.alignment = WD_ALIGN_PARAGRAPH.RIGHT; runner = p_empresa.add_run("CONSTRUCCIONES V & G LIZ CASTILLO GARCIA SPA\nRUT: 77.858.577-4\nDirección: Vilaco 301, Toconao\nTeléfono: +56 9 52341652"); runner.font.size = Pt(8); runner.bold = True; p_empresa.paragraph_format.space_after = Pt(0)
    document.add_paragraph("---" * 12).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_orden_info = document.add_paragraph(); p_orden_info.add_run(f"Orden de Compra #{orden.numero_venta or orden.id}\n").bold = True; p_orden_info.add_run(f"Cliente: {orden.cliente}\n"); p_orden_info.add_run(f"Rut: {orden.rut or 'N/A'}\n"); p_orden_info.add_run(f"Fecha: {orden.fecha.strftime('%d-%m-%Y %H:%M')}\n"); p_orden_info.add_run(f"Dirección: {orden.direccion or 'N/A'}"); [run.font.size for run in p_orden_info.runs]; p_orden_info.paragraph_format.space_after = Pt(6)
    document.add_paragraph("---" * 12).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_detalle_titulo = document.add_paragraph(); p_detalle_titulo.add_run("Detalle").bold = True; p_detalle_titulo.runs[0].font.size = Pt(10); p_detalle_titulo.paragraph_format.space_after = Pt(3)
    table = document.add_table(rows=1, cols=4); table.style = 'Table Grid'; table.autofit = False
    hdr_cells = table.rows[0].cells; hdr_cells[0].text = 'Producto'; hdr_cells[1].text = 'Cant'; hdr_cells[2].text = 'P. Unit.'; hdr_cells[3].text = 'Total'; [cell.paragraphs[0].runs[0].font for cell in hdr_cells]; [cell.paragraphs[0].alignment for cell in hdr_cells]
    for detalle in orden.detalles.all():
        row_cells = table.add_row().cells; precio_unit_str = f"${detalle.precio_unitario:,.0f}".replace(",","."); total_linea_str = f"${detalle.total_linea:,.0f}".replace(",","."); row_cells[0].text = detalle.producto.nombre; row_cells[1].text = str(detalle.cantidad); row_cells[2].text = precio_unit_str; row_cells[3].text = total_linea_str
        for i, cell in enumerate(row_cells): cell.paragraphs[0].runs[0].font.size = Pt(9); (cell.paragraphs[0].alignment) if i > 0 else None
    table.columns[0].width = Inches(2.8); table.columns[1].width = Inches(0.5); table.columns[2].width = Inches(0.9); table.columns[3].width = Inches(1.0)
    document.add_paragraph("---" * 12).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_total = document.add_paragraph(); p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT; total_str = f"${orden.total:,.0f}".replace(",","."); runner_total = p_total.add_run(f"Total a Pagar: {total_str}"); runner_total.bold = True; runner_total.font.size = Pt(11); p_total.paragraph_format.space_after = Pt(0)
    document.add_paragraph("---" * 12).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_gracias = document.add_paragraph("¡Gracias por su compra!\nEsperamos atenderle pronto."); p_gracias.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_gracias.runs[0].font.size = Pt(8)
    f = BytesIO(); document.save(f); f.seek(0)
    response = HttpResponse(f.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'); response['Content-Disposition'] = f'attachment; filename="orden_{orden.numero_venta or orden.id}.docx"'; f.close()
    return response

# --- Vistas de Trabajadores (CRUD) ---

@login_required
def lista_trabajadores(request):
    trabajadores = Trabajador.objects.all()
    return render(request, 'app/lista_trabajadores.html', {'trabajadores': trabajadores})

@login_required
def crear_trabajador(request):
    if request.method == 'POST':
        form = TrabajadorForm(request.POST)
        if form.is_valid(): form.save(); messages.success(request, "Trabajador creado exitosamente."); return redirect('lista_trabajadores')
    else: form = TrabajadorForm()
    return render(request, 'app/crear_trabajador.html', {'form': form})

@login_required
def editar_trabajador(request, pk):
    trabajador = get_object_or_404(Trabajador, pk=pk)
    if request.method == 'POST':
        form = TrabajadorForm(request.POST, instance=trabajador)
        if form.is_valid(): form.save(); messages.success(request, "Trabajador actualizado exitosamente."); return redirect('lista_trabajadores')
    else: form = TrabajadorForm(instance=trabajador)
    return render(request, 'app/editar_trabajador.html', {'form': form, 'trabajador': trabajador})

@login_required
def eliminar_trabajador(request, pk):
    trabajador = get_object_or_404(Trabajador, pk=pk)
    if request.method == 'POST': nombre_trabajador = trabajador.nombre; trabajador.delete(); messages.success(request, f"Trabajador '{nombre_trabajador}' eliminado exitosamente."); return redirect('lista_trabajadores')
    return render(request, 'app/confirmar_eliminar.html', {'object': trabajador, 'cancel_url': reverse('lista_trabajadores')})

# --- Vistas de Asistencia y Salarios ---

@login_required
def asistencia_manual(request):
    if request.method == 'POST':
        form = AsistenciaManualForm(request.POST)
        if form.is_valid():
            trabajador = form.cleaned_data['trabajador']; fecha = form.cleaned_data['fecha']; tipo_proyecto = form.cleaned_data['tipo_proyecto']
            if not Asistencia.objects.filter(trabajador=trabajador, fecha=fecha, tipo_proyecto=tipo_proyecto).exists(): Asistencia.objects.create(trabajador=trabajador, fecha=fecha, tipo_proyecto=tipo_proyecto); messages.success(request, "Asistencia registrada exitosamente."); return redirect('asistencia_confirmacion')
            else: messages.error(request, "Esta asistencia ya fue registrada.")
    else: form = AsistenciaManualForm()
    return render(request, 'app/asistencia_manual.html', {'form': form})

@login_required
def asistencia_confirmacion(request):
    return render(request, 'app/asistencia_confirmacion.html')

@login_required
def calcular_salario(request):
    context = {}
    if request.method == 'POST':
        form = CalculoSalarioForm(request.POST)
        if form.is_valid():
            trabajador = form.cleaned_data['trabajador']; fecha_inicio = form.cleaned_data['fecha_inicio']; fecha_fin = form.cleaned_data['fecha_fin']; tipo_proyecto = form.cleaned_data['tipo_proyecto']
            if trabajador.salario_por_dia <= 0: messages.error(request, f"El trabajador {trabajador.nombre} no tiene un salario por día definido.")
            else:
                asistencias_qs = Asistencia.objects.filter(trabajador=trabajador, fecha__range=[fecha_inicio, fecha_fin], tipo_proyecto=tipo_proyecto); asistencias_count = asistencias_qs.count(); salario_total = asistencias_count * trabajador.salario_por_dia
                context['mensaje'] = f"Cálculo para {trabajador.nombre} ({fecha_inicio.strftime('%d/%m/%Y')} - {fecha_fin.strftime('%d/%m/%Y')})"; context['asistencias'] = asistencias_count; context['salario_total'] = salario_total
                context['form_data'] = {'trabajador': trabajador.id, 'fecha_inicio': fecha_inicio.isoformat(), 'fecha_fin': fecha_fin.isoformat(), 'tipo_proyecto': tipo_proyecto, 'salario_total': float(salario_total)}; context['form_instance'] = form
        else: context['form_instance'] = form
    if 'form_instance' not in context: context['form'] = CalculoSalarioForm()
    else: context['form'] = context['form_instance']
    return render(request, 'app/calcular_salario.html', context)

@login_required
def registrar_pago_gasto(request):
    if request.method == 'POST':
        trabajador_id = request.POST.get('trabajador'); fecha_inicio_str = request.POST.get('fecha_inicio'); fecha_fin_str = request.POST.get('fecha_fin'); tipo_proyecto = request.POST.get('tipo_proyecto'); salario_total_str = request.POST.get('salario_total')
        if not all([trabajador_id, fecha_inicio_str, fecha_fin_str, tipo_proyecto, salario_total_str]): messages.error(request, "Faltan datos."); return redirect('calcular_salario')
        try:
            trabajador = Trabajador.objects.get(pk=trabajador_id); salario_total = float(salario_total_str)
            Gasto.objects.create(fecha=date.today(), categoria='SALARIO', descripcion=f"Pago salario a {trabajador.nombre} por período {fecha_inicio_str} al {fecha_fin_str} ({tipo_proyecto})", monto=salario_total, tipo_proyecto=tipo_proyecto)
            messages.success(request, "Salario registrado como Gasto."); return redirect('lista_gastos')
        except Trabajador.DoesNotExist: messages.error(request, "Trabajador no existe.")
        except ValueError: messages.error(request, "Monto inválido.")
        except Exception as e: messages.error(request, f"Error: {e}")
        return redirect('calcular_salario')
    return redirect('calcular_salario')

# --- Vistas de Inventario (CRUD) ---

@login_required
def inventario(request):
    productos = Producto.objects.all().order_by('nombre')
    return render(request, 'app/inventario.html', {'productos': productos})

@login_required
def crear_producto(request):
    titulo = "Agregar Nuevo Producto"
    if request.method == 'POST':
        form = ProductoForm(request.POST)
        if form.is_valid(): form.save(); messages.success(request, "Producto creado."); return redirect('inventario')
    else: form = ProductoForm()
    return render(request, 'app/crear_producto.html', {'form': form, 'titulo': titulo})

@login_required
def editar_producto(request, pk):
    producto = get_object_or_404(Producto, pk=pk); titulo = f"Editar Producto: {producto.nombre}"
    if request.method == 'POST':
        form = ProductoForm(request.POST, instance=producto)
        if form.is_valid(): form.save(); messages.success(request, "Producto actualizado."); return redirect('inventario')
    else: form = ProductoForm(instance=producto)
    return render(request, 'app/editar_producto.html', {'form': form, 'producto': producto, 'titulo': titulo})

@login_required
def eliminar_producto(request, pk):
    producto = get_object_or_404(Producto, pk=pk)
    if request.method == 'POST': nombre_producto = producto.nombre; producto.delete(); messages.success(request, f"Producto '{nombre_producto}' eliminado."); return redirect('inventario')
    return render(request, 'app/confirmar_eliminar.html', {'object': producto, 'cancel_url': reverse('inventario')})

# --- Vistas de Gastos (CRUD) ---

@login_required
def lista_gastos(request):
    gastos = Gasto.objects.all().order_by('-fecha')
    return render(request, 'app/lista_gastos.html', {'gastos': gastos})

@login_required
def registrar_gasto(request):
    if request.method == 'POST':
        form = GastoForm(request.POST)
        if form.is_valid(): form.save(); messages.success(request, "Gasto registrado."); return redirect('lista_gastos')
    else: form = GastoForm()
    return render(request, 'app/registrar_gasto.html', {'form': form})

@login_required
def editar_gasto(request, pk):
    gasto = get_object_or_404(Gasto, pk=pk)
    if request.method == 'POST':
        form = GastoForm(request.POST, instance=gasto)
        if form.is_valid(): form.save(); messages.success(request, "Gasto actualizado."); return redirect('lista_gastos')
    else: form = GastoForm(instance=gasto)
    return render(request, 'app/editar_gasto.html', {'form': form, 'gasto': gasto})

@login_required
def eliminar_gasto(request, pk):
    gasto = get_object_or_404(Gasto, pk=pk)
    if request.method == 'POST': gasto.delete(); messages.success(request, f"Gasto del {gasto.fecha} eliminado."); return redirect('lista_gastos')
    return render(request, 'app/confirmar_eliminar.html', {'object': gasto, 'cancel_url': reverse('lista_gastos')})

# --- API para JavaScript ---
@login_required
def get_stock_producto(request, producto_id):
    try: producto = Producto.objects.get(pk=producto_id); return JsonResponse({'stock': producto.stock})
    except Producto.DoesNotExist: return JsonResponse({'error': 'Producto no encontrado'}, status=404)
    except Exception as e: return JsonResponse({'error': str(e)}, status=500)

# --- NUEVAS VISTAS DE CONFIGURACIÓN ---

@login_required
def user_settings(request):
    """Página principal de configuración del usuario."""
    return render(request, 'app/user_settings.html')

@login_required
def edit_profile(request):
    """Vista para editar la información del usuario."""
    if request.method == 'POST':
        form = EditProfileForm(request.POST, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, '¡Tu perfil ha sido actualizado!')
            return redirect('user_settings') # Volver a la página de configuración
    else:
        # Pasamos la instancia del usuario actual al formulario
        form = EditProfileForm(instance=request.user)
    
    return render(request, 'app/edit_profile.html', {'form': form})

# Usaremos la vista incorporada de Django para cambiar contraseña,
# pero podemos personalizar la URL y las plantillas si queremos.
# Aquí definimos una subclase para especificar dónde redirigir después del éxito.
class CustomPasswordChangeView(PasswordChangeView):
    template_name='app/password_change_form.html' # Plantilla personalizada
    success_url = reverse_lazy('password_change_done') # Redirigir a la página de éxito

# No necesitamos una vista personalizada para 'done', Django la maneja bien,
# solo especificamos la plantilla en urls.py
# --- FIN DE NUEVAS VISTAS ---